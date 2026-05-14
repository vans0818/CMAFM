"""브리핑 시나리오 워드 문서 생성"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── 스타일 헬퍼 ──────────────────────────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_space(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    para_space(p, before=12, after=4)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=15, bold=True, color=(0x1F, 0x49, 0x7D))
    elif level == 2:
        set_font(run, size=12, bold=True, color=(0x2E, 0x74, 0xB5))
    else:
        set_font(run, size=11, bold=True, color=(0x40, 0x40, 0x40))
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    para_space(p, before=2, after=2, line=16)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_quote(doc, text):
    """브리핑 발화 스크립트 박스"""
    p = doc.add_paragraph()
    para_space(p, before=4, after=4, line=17)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    # 배경색 흉내 — 왼쪽 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "2E74B5")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, size=10.5, color=(0x1A, 0x1A, 0x5E))
    run.font.italic = True
    return p

def add_demo_box(doc, text):
    """시연 박스"""
    p = doc.add_paragraph()
    para_space(p, before=4, after=4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("  🖥️  " + text + "  ")
    set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
    # 배경색은 직접 지원 안 되므로 shading으로
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "2E74B5")
    rPr.append(shd)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        # 헤더 배경
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F497D")
        tcPr.append(shd)

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bold_cell = (c_idx == 0 and "CMAFM" in str(val)) or \
                        ("CMAFM" in str(row_data[0]))
            run = p.add_run(str(val))
            set_font(run, size=10, bold=bold_cell,
                     color=(0x1F, 0x49, 0x7D) if bold_cell else None)
            # 강조 행 배경
            if "CMAFM" in str(row_data[0]):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "DEEAF1")
                tcPr.append(shd)

    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])

    return table

def add_separator(doc):
    p = doc.add_paragraph()
    para_space(p, before=2, after=2)
    run = p.add_run("─" * 68)
    set_font(run, size=8, color=(0xBB, 0xBB, 0xBB))
    return p

# ══════════════════════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_space(p, before=40, after=6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CMAFM 브리핑 시나리오")
set_font(run, size=22, bold=True, color=(0x1F, 0x49, 0x7D))

p = doc.add_paragraph()
para_space(p, before=0, after=4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("교차 모달 어텐션 기반 RGB-Thermal 융합 객체 탐지")
set_font(run, size=14, color=(0x2E, 0x74, 0xB5))

p = doc.add_paragraph()
para_space(p, before=0, after=60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cross-Modal Attention Fusion Model for Multispectral Object Detection")
set_font(run, size=11, color=(0x70, 0x70, 0x70))
run.font.italic = True

p = doc.add_paragraph()
para_space(p, before=0, after=4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("총 브리핑 시간 : 5분  |  시연 포함")
set_font(run, size=11, bold=True)

p = doc.add_paragraph()
para_space(p, before=2, after=2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("M3FD 데이터셋 샘플 대시보드 실시간 시연 예정")
set_font(run, size=10, color=(0x60, 0x60, 0x60))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 브리핑 개요
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "브리핑 개요", level=1)

overview_table = doc.add_table(rows=5, cols=2)
overview_table.style = "Table Grid"
overview_data = [
    ("주제",    "교차 모달 어텐션 기반 RGB-Thermal 융합 객체 탐지 (CMAFM)"),
    ("총 시간",  "5분 (발표 약 4분 + 시연 약 1분)"),
    ("데이터셋", "M3FD — 4,200쌍 RGB+Thermal 도로 영상, 6개 클래스"),
    ("시연 도구", "CMAFM Detection Dashboard (Streamlit, localhost:8501)"),
    ("핵심 메시지", "야간·악천후 환경에서 RGB 단독 대비 mAP@0.5 +10.5%p 향상"),
]
for i, (k, v) in enumerate(overview_data):
    row = overview_table.rows[i]
    kc = row.cells[0]
    vc = row.cells[1]
    kc.width = Cm(3.5); vc.width = Cm(12.5)
    kp = kc.paragraphs[0]; vp = vc.paragraphs[0]
    kr = kp.add_run(k); vr = vp.add_run(v)
    set_font(kr, size=10, bold=True, color=(0x1F, 0x49, 0x7D))
    set_font(vr, size=10)
    tcPr = kc._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DEEAF1")
    tcPr.append(shd)

doc.add_paragraph()

# ── 타임라인 표 ──
add_heading(doc, "타임라인 요약", level=2)
add_table(doc,
    headers=["시간", "구간", "내용", "시연"],
    rows=[
        ["0:00 – 0:40", "오프닝",       "문제 제기 — RGB·Thermal 각각의 한계",       ""],
        ["0:40 – 1:40", "모델 구조",    "CMAFM 아키텍처 및 교차 어텐션 원리",         ""],
        ["1:40 – 2:30", "데이터·시연①", "M3FD 데이터셋 소개 + 주간/야간 샘플 탐지",  "📷 샘플 테스트"],
        ["2:30 – 3:30", "성능 결과",    "5개 모델 비교 — 수치로 증명",               ""],
        ["3:30 – 4:10", "Ablation·시연②","구성 요소 기여도 + 이미지 직접 업로드",    "🖼️ 이미지 탐지"],
        ["4:10 – 5:00", "결론·확장",    "정리 및 향후 발전 방향 3가지",              ""],
    ],
    col_widths=[2.8, 2.5, 7.5, 3.2]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 섹션 1
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "① 오프닝 — 문제 제기  [0:00 – 0:40]", level=1)
add_separator(doc)

add_heading(doc, "발화 스크립트", level=2)
add_quote(doc,
    "자율주행, 군사 감시, 보안 시스템에서 객체 탐지는 핵심 기술입니다.\n"
    "그런데 기존 RGB 카메라 기반 탐지는 야간, 역광, 안개 환경에서 성능이 급격히 떨어집니다.\n"
    "반대로 Thermal 카메라는 조명에 독립적이지만 텍스처와 색상 정보가 없어\n"
    "유사한 열원을 구분하기 어렵습니다.\n"
    "이 두 카메라의 약점을 서로 보완하면 어떨까요?\n"
    "오늘은 RGB와 Thermal 영상을 딥러닝으로 융합하여 더 강인한 객체 탐지를 실현한\n"
    "CMAFM 모델을 소개합니다."
)

add_heading(doc, "핵심 포인트", level=2)
for pt in [
    "RGB 카메라: 주간·텍스처 강점 / 야간·역광·안개에서 성능 급락",
    "Thermal 카메라: 조명 독립적 / 텍스처·색상 부재로 유사 열원 구분 어려움",
    "두 모달리티의 상호 보완 = 모든 환경에서 강인한 탐지",
]:
    p = doc.add_paragraph(style="List Bullet")
    para_space(p, before=1, after=1)
    run = p.add_run(pt)
    set_font(run, size=10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 섹션 2
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "② 모델 구조 설명  [0:40 – 1:40]", level=1)
add_separator(doc)

add_heading(doc, "발화 스크립트", level=2)
add_quote(doc,
    "모델은 크게 4단계로 구성됩니다.\n"
    "RGB와 Thermal 영상 각각에 독립적인 ResNet-50 백본을 배치하고,\n"
    "두 백본이 추출한 특징맵을 CMAFM — 교차 모달 어텐션 융합 모듈 — 에 통과시킵니다.\n"
    "이후 FPN으로 다중 스케일 특징을 생성하고, Faster R-CNN 헤드가 최종 탐지를 수행합니다.\n\n"
    "핵심은 CMAFM입니다.\n"
    "기존 방식은 두 영상을 단순히 붙이거나 더하는 수준이었습니다.\n"
    "CMAFM은 RGB가 Thermal을 참조하고, Thermal이 RGB를 참조하며\n"
    "서로 부족한 정보를 채워주는 양방향 어텐션을 수행합니다.\n"
    "예를 들어 야간에 RGB가 어두워 놓친 사람을, Thermal의 열 정보로 되살려 탐지합니다.\n\n"
    "또한 Transformer 방식의 O(N²) 연산 복잡도 문제를 피하기 위해\n"
    "GAP 기반 채널 어텐션으로 O(C) 복잡도를 유지해 고해상도에서도 효율적입니다."
)

add_heading(doc, "모델 구조 다이어그램", level=2)
p = doc.add_paragraph()
para_space(p, before=2, after=2)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "RGB 영상    →  ResNet-50 백본  ─┐\n"
    "                                 ├→  CMAFM  →  FPN  →  Faster R-CNN  →  탐지 결과\n"
    "Thermal 영상 →  ResNet-50 백본  ─┘"
)
set_font(run, name="Courier New", size=9.5)

add_heading(doc, "CMAFM 3단계", level=2)
add_table(doc,
    headers=["단계", "명칭", "동작", "복잡도"],
    rows=[
        ["①", "채널 교차 어텐션", "GAP으로 채널 통계 추출 → 교차 모달 dot-product 스케일링", "O(C)"],
        ["②", "공간 교차 게이팅", "DWConv로 공간 특징 추출 → 상대 모달리티 맥락으로 게이팅", "O(C·H·W)"],
        ["③", "게이트 융합",      "픽셀·채널별 게이트 α로 두 모달리티 적응적 가중 합산 + 잔차", "O(C·H·W)"],
    ],
    col_widths=[1.2, 3.8, 8.5, 2.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 섹션 3
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "③ 데이터셋 소개 + 시연①  [1:40 – 2:30]", level=1)
add_separator(doc)

add_heading(doc, "발화 스크립트", level=2)
add_quote(doc,
    "실험은 M3FD 데이터셋을 사용했습니다.\n"
    "도로 환경에서 수집한 4,200쌍의 RGB+Thermal 영상으로,\n"
    "사람, 차량, 버스, 오토바이, 가로등, 트럭 6개 클래스를 탐지합니다.\n"
    "주간 779장, 야간 61장으로 구성되어 있으며\n"
    "지금 보시는 대시보드에서 이 데이터로 바로 탐지를 시연해드리겠습니다."
)

add_demo_box(doc, "시연 ① — 대시보드 [📷 샘플 테스트] 탭")
add_body(doc, "① 슬라이더로 주간 샘플 선택 → [탐지 실행] 클릭", indent=True)
add_body(doc, '   → "이게 주간 샘플입니다. 차량과 사람이 정확하게 탐지되는 것을 볼 수 있습니다."', indent=True)
add_body(doc, "② 슬라이더로 야간 샘플 선택 → [탐지 실행] 클릭", indent=True)
add_body(doc, '   → "이건 야간입니다. RGB만으로는 어둠 속에서 보이지 않을 객체들이 Thermal 덕분에 잡힙니다."', indent=True)

add_heading(doc, "M3FD 데이터셋 클래스 구성", level=2)
add_table(doc,
    headers=["클래스", "객체 수", "비율"],
    rows=[
        ["Car",        "18,296", "53.2%"],
        ["People",     "11,477", "33.4%"],
        ["Lamp",        "2,405",  "7.0%"],
        ["Truck",       "1,008",  "2.9%"],
        ["Bus",           "700",  "2.0%"],
        ["Motorcycle",    "521",  "1.5%"],
        ["합계",        "34,407", "100%"],
    ],
    col_widths=[4.0, 5.0, 5.0]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 섹션 4
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "④ 성능 결과 — 수치로 증명  [2:30 – 3:30]", level=1)
add_separator(doc)

add_heading(doc, "발화 스크립트", level=2)
add_quote(doc,
    "성능 수치로 보겠습니다. 5가지 방법을 비교했습니다.\n\n"
    "RGB 단독 대비 +10.5%p, 단순 융합 대비 +8.7%p 향상입니다.\n\n"
    "특히 주목할 수치는 야간 성능입니다.\n"
    "RGB만 쓰면 야간에 62.8%로 주간과 거의 차이가 없습니다 — 이미 성능이 한계에 도달했기 때문입니다.\n"
    "반면 CMAFM은 야간에 85.3%로, 주간(72.8%) 대비 오히려 12.5%p 더 높습니다.\n"
    "어두울수록 Thermal이 RGB를 더 강하게 보완하기 때문입니다."
)

add_heading(doc, "모델별 성능 비교", level=2)
add_table(doc,
    headers=["방법", "mAP@0.5 (전체)", "mAP@0.5 (야간)", "mAP@0.5 (주간)", "Recall"],
    rows=[
        ["Thermal만 사용",          "52.8%", "68.6%", "51.5%", "72.2%"],
        ["RGB만 사용",              "63.2%", "62.8%", "63.0%", "82.9%"],
        ["Early Fusion (단순 융합)", "65.0%", "80.0%", "64.2%", "83.5%"],
        ["이중 백본 (어텐션 없음)",  "70.0%", "79.9%", "69.2%", "86.7%"],
        ["CMAFM (제안 모델)",       "73.7%", "85.3%", "72.8%", "87.4%"],
    ],
    col_widths=[4.5, 3.2, 3.2, 3.2, 2.4]
)

add_heading(doc, "야간 성능이 주간보다 높은 이유", level=2)
for pt in [
    "야간 → RGB 카메라 노이즈 증가, 텍스처 손실 → RGB 단독 성능 한계 도달",
    "야간 → Thermal 카메라는 조명 무관 → 열원(사람·차량) 더 선명하게 포착",
    "두 모달리티 중 Thermal이 지배적 기여 → 융합 효과 극대화",
    "버스 야간 Recall 100%, 차량 야간 Recall 98.9% 달성",
]:
    p = doc.add_paragraph(style="List Bullet")
    para_space(p, before=1, after=1)
    run = p.add_run(pt)
    set_font(run, size=10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 섹션 5
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "⑤ Ablation Study + 시연②  [3:30 – 4:10]", level=1)
add_separator(doc)

add_heading(doc, "발화 스크립트", level=2)
add_quote(doc,
    "어떤 구성 요소가 얼마나 기여했는지 분해해봤습니다.\n\n"
    "이중 백본이 5.0%p, 교차 어텐션이 추가로 3.7%p를 담당합니다.\n"
    "RGB와 Thermal은 통계적 특성이 완전히 다르기 때문에,\n"
    "하나의 백본으로 두 영상을 동시에 처리하는 것보다\n"
    "각자 독립 백본을 가지는 것이 유리하다는 것이 수치로 검증됩니다."
)

add_heading(doc, "구성 요소별 기여도", level=2)
add_table(doc,
    headers=["구성", "mAP@0.5", "이전 대비 향상"],
    rows=[
        ["Early Fusion (베이스라인)", "65.0%", "기준"],
        ["+ 이중 백본 구조",          "70.0%", "+5.0%p"],
        ["+ 교차 모달 어텐션 (CMAFM)", "73.7%", "+3.7%p"],
    ],
    col_widths=[6.5, 4.0, 5.5]
)

add_demo_box(doc, "시연 ② — 대시보드 [🖼️ 이미지 탐지] 탭")
add_body(doc, "① 'RGB만 업로드' 모드 선택", indent=True)
add_body(doc, "② M3FD 야간 샘플 이미지 업로드 → [탐지 실행]", indent=True)
add_body(doc, "③ 사이드바 신뢰도 임계값 슬라이더 조작 → 탐지 민감도 실시간 변화 시연", indent=True)
add_body(doc, '   → "신뢰도 임계값을 조절하면서 탐지 민감도를 실시간으로 바꿀 수 있습니다."', indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 섹션 6
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "⑥ 결론 및 확장 가능성  [4:10 – 5:00]", level=1)
add_separator(doc)

add_heading(doc, "발화 스크립트", level=2)
add_quote(doc,
    "정리하겠습니다.\n"
    "CMAFM은 RGB와 Thermal이 서로를 참조하는 교차 어텐션을 통해\n"
    "단일 모달리티의 한계를 극복하고, 특히 야간·악천후 환경에서 강인한 탐지를 실현했습니다.\n\n"
    "향후 발전 방향은 세 가지입니다.\n\n"
    "첫째, 실시간 하드웨어 결합 — RGB+Thermal 듀얼 카메라를 물리적으로 결합하고\n"
    "캘리브레이션하면 현재 대시보드에 실시간 카메라 입력으로 바로 연결 가능합니다.\n\n"
    "둘째, 경량화 — ResNet-50을 MobileNet 계열로 교체하면\n"
    "드론, 차량용 ECU 등 엣지 디바이스에 탑재 가능합니다.\n\n"
    "셋째, 도메인 확장 — 현재 도로 환경 데이터로 학습했지만,\n"
    "의료 열화상, 산불 감지, 군사 감시 영역으로 확장할 수 있습니다.\n\n"
    "현재 상용 제품 대부분이 탐지 결과를 단순 합산하는 수준인 반면,\n"
    "CMAFM은 특징 레벨에서 융합하는 학계 최신 방향입니다.\n"
    "이 연구가 열악한 환경에서도 신뢰할 수 있는 객체 탐지 시스템의 기반이 되길 기대합니다.\n"
    "감사합니다."
)

add_heading(doc, "향후 발전 방향", level=2)
add_table(doc,
    headers=["방향", "내용", "기대 효과"],
    rows=[
        ["① 실시간 하드웨어 결합",
         "RGB+Thermal 듀얼 카메라 물리 결합\n+ 캘리브레이션",
         "파일 업로드 없이 실시간\n카메라 스트림으로 탐지"],
        ["② 경량화",
         "ResNet-50 → MobileNet/EfficientNet\n백본 교체",
         "드론·차량용 ECU 등\n엣지 디바이스 탑재"],
        ["③ 도메인 확장",
         "의료 열화상 / 산불 감지 /\n군사 감시 데이터 학습",
         "다양한 산업 분야 적용"],
    ],
    col_widths=[4.0, 6.5, 5.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 부록 — 핵심 수치 요약
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "부록 — 핵심 수치 한눈에 보기", level=1)
add_separator(doc)

add_heading(doc, "최종 모델 성능 (CMAFM, best.pth 기준)", level=2)
add_table(doc,
    headers=["조건", "mAP@0.5", "mAP@[.5:.95]", "Recall", "F1", "Miss Rate", "FPS"],
    rows=[
        ["전체", "73.7%", "40.6%", "87.4%", "79.9%", "12.6%", "11.9"],
        ["야간", "85.3%", "50.7%", "90.9%", "88.0%",  "9.1%", "11.3"],
        ["주간", "72.8%", "39.8%", "87.0%", "79.3%", "13.0%",  "6.6"],
    ],
    col_widths=[2.2, 2.5, 3.2, 2.5, 2.0, 2.8, 1.8]
)

add_heading(doc, "클래스별 성능 (전체 조건)", level=2)
add_table(doc,
    headers=["클래스", "AP@0.5", "Recall", "Miss Rate"],
    rows=[
        ["People",     "75.7%", "87.5%", "12.5%"],
        ["Car",        "87.3%", "93.6%",  "6.4%"],
        ["Bus",        "77.3%", "89.8%", "10.2%"],
        ["Motorcycle", "69.7%", "81.2%", "18.8%"],
        ["Lamp",       "67.4%", "84.6%", "15.4%"],
        ["Truck",      "64.9%", "87.7%", "12.3%"],
    ],
    col_widths=[4.0, 4.0, 4.0, 4.0]
)

add_heading(doc, "시연 체크리스트", level=2)
for item in [
    "대시보드 실행 확인 (run_dashboard.bat → http://localhost:8501)",
    "모델 로드 완료 확인 (runs/best.pth → '✅ 모델 로드 완료!')",
    "샘플 테스트 탭 — 주간 샘플 탐지 준비 (슬라이더 인덱스 0~50 범위)",
    "샘플 테스트 탭 — 야간 샘플 탐지 준비 (밝기 낮은 샘플 미리 확인)",
    "이미지 탐지 탭 — 야간 샘플 이미지 파일 별도 저장 준비",
    "신뢰도 임계값 기본값 0.5 확인",
]:
    p = doc.add_paragraph(style="List Bullet")
    para_space(p, before=1, after=1)
    run = p.add_run(item)
    set_font(run, size=10.5)

# ── 저장 ──────────────────────────────────────────────────────────────────────
out_path = r"d:\★RGB-LWIR(멘토ver-최종)\CMAFM_브리핑_시나리오.docx"
doc.save(out_path)
print(f"저장 완료: {out_path}")
