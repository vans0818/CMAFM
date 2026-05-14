"""
학회지 2단 편집 Word 문서 생성기
- 제목, 요약/Abstract: 1단 (전체 폭)
- I장 이후 본문: 2단 편집
- 표: 2단 폭에 맞게 축소
- 그림: 2단 폭(컬럼 폭) 기준으로 크기 조정
"""

import os
os.environ["PYTHONUTF8"] = "1"

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE   = Path(__file__).parent
OUT    = BASE / "runs" / "학회지_2단편집.docx"
IMGS   = BASE / "runs" / "paper_figures"
ARCH   = BASE.parent / "image" / "학회지(수정)"

# ── 치수 상수 ─────────────────────────────────────────────────────────────────
PAGE_W      = Cm(21.0)
MARGIN_LR   = Cm(2.5)
MARGIN_TB   = Cm(2.5)
COL_SPACE   = Cm(1.0)
BODY_W      = PAGE_W - MARGIN_LR * 2                        # 16 cm
COL_W       = (BODY_W - COL_SPACE) / 2                      # ~7.5 cm
FULL_W      = BODY_W                                         # 표/그림 전체폭용


# ── XML 헬퍼 ─────────────────────────────────────────────────────────────────
def set_2col(section):
    """섹션을 2단으로 설정."""
    sectPr = section._sectPr
    # 기존 cols 제거
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "720")   # 0.5인치 = 720 twips ≈ 1.27cm (조정)
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def set_1col(section):
    """섹션을 1단으로 설정."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    sectPr.append(cols)


def add_section_break(doc, break_type="continuous"):
    """섹션 구분자 삽입. break_type: continuous | nextPage"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), str(int(PAGE_W.pt * 20)))
    pgSz.set(qn("w:h"), str(int(Cm(29.7).pt * 20)))
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"),    str(int(MARGIN_TB.pt * 20)))
    pgMar.set(qn("w:bottom"), str(int(MARGIN_TB.pt * 20)))
    pgMar.set(qn("w:left"),   str(int(MARGIN_LR.pt * 20)))
    pgMar.set(qn("w:right"),  str(int(MARGIN_LR.pt * 20)))
    sectPr.append(pgMar)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    sectPr.append(cols)
    # 연속 구분
    if break_type == "continuous":
        typ = OxmlElement("w:type")
        typ.set(qn("w:val"), "continuous")
        sectPr.append(typ)
    pPr.append(sectPr)
    return p


# ── 스타일 헬퍼 ──────────────────────────────────────────────────────────────
def p_title(doc, text, size=16, bold=True, center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))
    return p


def p_body(doc, text, size=9, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))
    return p


def p_heading(doc, text, level=1, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))
    return p


def p_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = "맑은 고딕"
    run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))
    return p


def _eastAsia(name):
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    return rFonts


def add_image(doc, img_path, width):
    """이미지 삽입. width: Cm or Inches 객체."""
    if not Path(img_path).exists():
        p_body(doc, f"[이미지 없음: {Path(img_path).name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run()
    run.add_picture(str(img_path), width=width)


# ── 표 스타일 ────────────────────────────────────────────────────────────────
def style_table(table, font_size=8):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "맑은 고딕"
                    run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))


def set_table_width(table, width_emu):
    """width_emu: EMU 단위 (docx.shared.Cm() 반환값)"""
    width_cm = width_emu / 914400 * 2.54  # EMU → inch → cm
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))  # 1cm ≈ 567 twips
    tblW.set(qn("w:type"), "dxa")
    old = tblPr.find(qn("w:tblW"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblW)


# ── 본문 생성 ────────────────────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # 페이지 설정 (기본 섹션 = 1단)
    section = doc.sections[0]
    section.page_width  = PAGE_W
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = MARGIN_LR
    section.top_margin  = section.bottom_margin = MARGIN_TB
    set_1col(section)

    # ══════════════════════════════════════════════════════
    # [1단 영역] 제목 + 요약/Abstract
    # ══════════════════════════════════════════════════════
    p_title(doc, "교차 모달 어텐션 기반 RGB-Thermal 영상 융합 객체 검출", size=15)
    p_title(doc, "Cross-Modal Attention-Based RGB-Thermal Image Fusion for Object Detection", size=11, bold=False)
    doc.add_paragraph()

    p_heading(doc, "요약", size=10)
    p_body(doc,
        "본 논문은 가시광(RGB) 영상과 적외선(Thermal) 영상을 효과적으로 융합하여 보행자 및 차량을 검출하는 "
        "딥러닝 기반 다중 스펙트럼 객체 검출 모델을 제안한다. 제안하는 모델은 각 모달리티에 독립적인 ResNet-50 "
        "백본을 할당하여 고유 특징을 추출한 후, 채널 교차 어텐션(Channel Cross-Attention)과 공간 교차 게이팅"
        "(Spatial Cross-Gating)으로 구성된 교차 모달 어텐션 융합 모듈(CMAFM)을 통해 상호 보완적 정보를 교환한다. "
        "M3FD 데이터셋 실험에서 mAP@0.5 73.7%, Recall 87.4%를 달성하였으며, RGB 단일 모달리티 대비 mAP@0.5 "
        "10.5%p, Recall 4.5%p 향상되었다. 야간 조건에서 mAP@0.5 85.3%, Recall 90.9%로 주간 대비 각각 "
        "12.5%p, 3.9%p 높은 성능을 보여 열악한 조도 환경에서의 효과를 실증하였다.",
        indent=True)
    p_body(doc, "핵심어: 다중 스펙트럼 객체 검출, RGB-Thermal 융합, 교차 모달 어텐션, 이중 백본, Faster R-CNN",
           bold=False, size=8)
    doc.add_paragraph()

    p_heading(doc, "Abstract", size=10)
    p_body(doc,
        "This paper proposes a deep learning-based multispectral object detection model that effectively fuses "
        "visible (RGB) and infrared (Thermal) images for pedestrian and vehicle detection. The proposed model "
        "assigns independent ResNet-50 backbones to each modality, then exchanges complementary information "
        "through a Cross-Modal Attention Fusion Module (CMAFM). Experiments on the M3FD dataset achieve "
        "73.7% mAP@0.5 and 87.4% Recall, with improvements of 10.5%p and 4.5%p over the RGB-only baseline. "
        "Under nighttime conditions, the model achieves 85.3% mAP@0.5 and 90.9% Recall, surpassing daytime "
        "performance by 12.5%p and 3.9%p respectively.",
        indent=True)
    p_body(doc, "Keywords: Multispectral Object Detection, RGB-Thermal Fusion, Cross-Modal Attention, Dual Backbone, Faster R-CNN",
           bold=False, size=8)

    # ── 연속 섹션 구분 → 2단 시작 ──────────────────────────────────────────
    add_section_break(doc, "continuous")

    # 새 섹션 추가 (2단)
    new_sec = doc.add_section()
    new_sec.page_width  = PAGE_W
    new_sec.page_height = Cm(29.7)
    new_sec.left_margin = new_sec.right_margin = MARGIN_LR
    new_sec.top_margin  = new_sec.bottom_margin = MARGIN_TB
    set_2col(new_sec)

    # ══════════════════════════════════════════════════════
    # [2단 영역] I. 서론
    # ══════════════════════════════════════════════════════
    p_heading(doc, "I. 서론", size=10)
    p_body(doc,
        "자율 주행, 군사 감시, 스마트 시티 등의 응용 분야에서 객체 검출(Object Detection)은 핵심 기술로 "
        "자리잡고 있다[1]. 특히 보행자 및 차량 검출은 안전과 직결되는 과제로, 다양한 환경 조건에서의 강인한 "
        "검출 성능이 요구된다.", indent=True)
    p_body(doc,
        "기존의 가시광(RGB) 카메라 기반 객체 검출은 주간 및 조명이 충분한 환경에서 우수한 성능을 보이나, "
        "야간, 안개, 역광 등 열악한 조건에서는 성능이 급격히 저하된다[2]. 반면, 적외선(Thermal) 카메라는 "
        "물체의 열 복사를 감지하므로 조명 조건에 독립적이나, 텍스처 및 색상 정보의 부재로 인해 유사한 열 "
        "특성을 가진 객체의 구별이 어렵다[3].", indent=True)
    p_body(doc,
        "이러한 한계를 극복하기 위해 RGB-Thermal(RGB-T) 융합 기반 객체 검출 연구가 활발히 진행되고 있다[4-6]. "
        "기존 융합 방법은 조기 융합(Early Fusion), 중간 수준 융합(Mid-level Fusion), 후기 융합(Late Fusion)으로 "
        "분류되며, 중간 수준 융합이 가장 높은 성능을 보이는 것으로 보고되고 있다[7]. 그러나 기존 방법들은 "
        "단순 연결(Concatenation)이나 요소별 덧셈을 사용하여 상호 보완적 관계를 충분히 활용하지 못한다.", indent=True)
    p_body(doc,
        "본 논문에서는 교차 모달 어텐션 융합 모듈(CMAFM)을 포함한 이중 백본 기반 다중 스펙트럼 객체 검출 "
        "모델을 제안한다. 주요 기여는 다음과 같다: (1) 메모리 효율적 CMAFM 제안, (2) 다중 스케일 이중 백본 "
        "아키텍처 설계, (3) M3FD 데이터셋에서의 포괄적 실험 및 ablation study.", indent=True)

    # ══════════════════════════════════════════════════════
    # III. 제안 방법
    # ══════════════════════════════════════════════════════
    p_heading(doc, "III. 제안 방법", size=10)

    p_heading(doc, "3.1 전체 구조", size=9)
    p_body(doc,
        "제안하는 모델은 (1) 이중 백본, (2) CMAFM, (3) FPN, (4) Faster R-CNN 검출 헤드로 구성된다. "
        "입력 RGB 및 Thermal 영상은 각각 640×640으로 리사이즈되며, 각 백본은 C3, C4, C5 세 개 스케일의 "
        "특징맵을 추출한다. 동일 스케일의 RGB-Thermal 특징쌍에 대해 CMAFM이 적용된다.", indent=True)

    arch_img = ARCH / "1776748944905.png"
    add_image(doc, arch_img, width=COL_W)
    p_caption(doc, "그림 1. 제안 모델의 전체 구조")

    p_heading(doc, "3.2 이중 백본 구조", size=9)
    p_body(doc,
        "각 모달리티에 독립적인 ResNet-50 백본을 할당한다. 두 백본은 동일한 ImageNet 사전 학습 가중치로 "
        "초기화되지만 학습 과정에서 독립적으로 미세 조정된다. RGB 백본은 텍스처·색상·형태 등 시각적 특징 "
        "추출에, Thermal 백본은 열 분포·열원 경계 등 적외선 고유 특징 추출에 특화된다.", indent=True)

    p_heading(doc, "3.3 교차 모달 어텐션 융합 모듈 (CMAFM)", size=9)
    p_body(doc, "CMAFM은 세 단계로 구성된다.", indent=True)

    p_heading(doc, "3.3.1 채널 교차 어텐션", size=9)
    p_body(doc,
        "전역 평균 풀링(GAP)으로 채널별 통계량을 추출한 후 교차 dot-product 스케일링을 수행한다. "
        "계산 복잡도가 O(C)로 매우 효율적이며 양방향 채널 정보 교환을 달성한다.", indent=True)

    p_heading(doc, "3.3.2 공간 교차 게이팅", size=9)
    p_body(doc,
        "DWConv 기반 공간 특징에 상대 모달리티의 맥락으로 게이팅을 수행하여 관련 영역을 강화하고 "
        "무관한 영역을 억제한다.", indent=True)

    p_heading(doc, "3.3.3 게이트 융합", size=9)
    p_body(doc,
        "학습 가능한 게이트 α(픽셀·채널별 독립)로 두 모달리티 특징을 가중 융합하고 잔차 연결을 적용한다: "
        "F_out = Conv₃ₓ₃(α·F_r'' + (1−α)·F_t'') + F_fused", indent=True)

    p_heading(doc, "3.4 검출 헤드", size=9)
    p_body(doc,
        "융합된 C3~C5 특징맵은 FPN(출력 256ch)을 거쳐 Faster R-CNN의 RPN이 앵커 기반 영역 제안을 생성하고, "
        "ROI Head가 분류 및 바운딩 박스 회귀를 수행한다.", indent=True)

    p_heading(doc, "3.5 학습 전략", size=9)
    p_body(doc,
        "백본에는 기본 학습률(0.005)의 0.1배 차등 학습률을 적용하여 사전 학습 가중치를 보존하고, "
        "융합 모듈 및 검출 헤드는 기본 학습률로 학습한다. "
        "StepLR 스케줄러로 10 epoch마다 학습률을 0.1배 감소시킨다.", indent=True)

    # ══════════════════════════════════════════════════════
    # II. 관련 연구
    # ══════════════════════════════════════════════════════
    p_heading(doc, "II. 관련 연구", size=10)

    p_heading(doc, "2.1 RGB-Thermal 융합 객체 검출", size=9)
    p_body(doc,
        "RGB-Thermal 융합 연구는 융합 위치에 따라 조기 융합, 중간 수준 융합, 후기 융합으로 분류된다. "
        "Liu et al.[4]의 TarDAL은 적대적 학습 기반 융합을 제안하였으며, Zhang et al.[5]는 반복적 "
        "융합-정제(Cyclic Fuse-and-Refine) 블록으로 두 모달리티의 특징을 점진적으로 통합하였다. "
        "Cao et al.[6]은 교차 모달 특징 융합을 Transformer 구조로 구현하였으나, 전체 공간 어텐션의 "
        "O(N²) 복잡도로 인해 고해상도 특징맵 적용에 제약이 있다.", indent=True)

    p_heading(doc, "2.2 어텐션 기반 특징 융합", size=9)
    p_body(doc,
        "Hu et al.[8]의 SENet은 채널 어텐션(Squeeze-and-Excitation)으로 채널별 중요도를 재조정하는 "
        "기법을 제안하였다. Zhang et al.[10]은 교차 모달리티 상호 어텐션 네트워크(CIAN)를 통해 두 "
        "스펙트럼 간 정보 교환을 시도하였다. 본 논문의 CMAFM은 채널 교차 어텐션과 합성곱 기반 공간 교차 "
        "게이팅을 결합하여 N×N 어텐션 행렬 없이 전역 및 지역 수준의 양방향 정보 교환을 달성한다는 점에서 "
        "기존 방법과 차별화된다.", indent=True)

    # ══════════════════════════════════════════════════════
    # IV. 실험
    # ══════════════════════════════════════════════════════
    p_heading(doc, "IV. 실험", size=10)

    p_heading(doc, "4.1 데이터셋", size=9)
    p_body(doc,
        "M3FD(Multi-Modal Multi-Scene Fusion Detection) 데이터셋[15]을 사용하였다. M3FD는 도로 주행 "
        "환경에서 수집된 4,200쌍의 정합된 가시광-적외선 영상(1024×768)으로 구성되며, 주간·야간·흐림·역광·"
        "안개 등 다양한 환경을 포함한다. 6개 클래스(People, Car, Bus, Motorcycle, Lamp, Truck)에 대해 "
        "총 34,407개의 바운딩 박스 어노테이션을 제공한다. 데이터셋을 8:2 비율로 학습(3,360쌍)/검증(840쌍)으로 "
        "분할하였다.", indent=True)

    # 표 1
    p_body(doc, "표 1. M3FD 데이터셋 클래스별 객체 수", bold=True, size=8)
    t1 = doc.add_table(rows=8, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t1, COL_W)
    data1 = [["클래스","객체 수","비율"],
             ["Car","18,296","53.2%"],["People","11,477","33.4%"],
             ["Lamp","2,405","7.0%"],["Truck","1,008","2.9%"],
             ["Bus","700","2.0%"],["Motorcycle","521","1.5%"],
             ["합계","34,407","100%"]]
    for i, row_data in enumerate(data1):
        row = t1.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    style_table(t1, font_size=8)
    for cell in t1.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    p_heading(doc, "4.2 실험 환경", size=9)
    p_body(doc,
        "실험 환경은 표 2와 같다. 데이터 증강으로 수평 반전(p=0.5)과 밝기/대비 변환(p=0.3)을 적용하였으며, "
        "동일한 변환이 RGB-Thermal 쌍에 함께 적용된다.", indent=True)

    # 표 2
    p_body(doc, "표 2. 실험 환경", bold=True, size=8)
    t2 = doc.add_table(rows=8, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t2, COL_W)
    data2 = [["항목","사양"],["GPU","NVIDIA RTX A6000 (48GB)"],
             ["Framework","PyTorch 2.4.0"],["Backbone","ResNet-50 (ImageNet pretrained)"],
             ["Optimizer","SGD (momentum=0.9, wd=5×10⁻⁴)"],
             ["학습률","0.005 (backbone: 0.0005)"],
             ["Batch / Epoch","8 / 30"],["Input size","640×640"]]
    for i, row_data in enumerate(data2):
        row = t2.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    style_table(t2, font_size=8)
    for cell in t2.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    p_heading(doc, "4.3 평가 지표", size=9)
    p_body(doc,
        "COCO 표준 평가 지표를 사용하였다. mAP@0.5는 IoU 임계값 0.5에서의 평균 정밀도이며, "
        "mAP@[.5:.95]는 IoU 0.5~0.95까지 0.05 간격 평균으로 더 엄격한 지표이다. "
        "Recall은 전체 정답 객체 대비 검출 비율, Miss Rate는 미검출 비율(1−Recall)이다.", indent=True)

    p_heading(doc, "4.4 Ablation Study", size=9)
    p_body(doc,
        "제안 모델의 각 구성 요소 기여도를 5가지 변형 모델로 검증하였다(표 3). "
        "평가 지표는 COCO 표준인 mAP@0.5 및 mAP@[.5:.95]를 사용하였다.", indent=True)

    # Ablation 이미지 (표 3 대체)
    abl_img = ARCH / "1776836203144.png"
    add_image(doc, abl_img, width=COL_W)
    p_caption(doc, "표 3. Ablation Study 결과")

    p_heading(doc, "4.4.1 모달리티 융합 효과", size=9)
    p_body(doc,
        "RGB 단일(0.632)에 Thermal 융합 시 mAP@0.5 +10.5%p, Recall +4.5%p(0.829→0.874)를 달성하였다. "
        "Thermal-only(0.528)는 야간(0.686)이 주간(0.515)보다 17.1%p 높아 LWIR의 조명 독립성이 확인된다. "
        "이러한 두 모달리티의 상호 보완적 특성이 융합 모델 야간 성능 향상(+12.5%p)의 근본 원인이다.", indent=True)

    p_heading(doc, "4.4.2 이중 백본 구조의 효과", size=9)
    p_body(doc,
        "Early Fusion(0.650) 대비 Dual+Concat(0.700)이 5.0%p 우수하였다. RGB와 Thermal의 통계적 특성 "
        "차이가 크므로 공유 백본의 동시 최적화가 어렵기 때문으로 해석된다.", indent=True)

    p_heading(doc, "4.4.3 교차 모달 어텐션의 효과", size=9)
    p_body(doc,
        "Dual+Concat(0.700) 대비 Full 모델(0.737)에서 CMAFM의 기여도는 +3.7%p로 확인되었다. "
        "채널 수준 전역 정보 교환과 공간 수준 지역 게이팅이 상호 보완적 특징 강화를 달성한다.", indent=True)

    p_heading(doc, "4.5 학습 곡선 분석", size=9)
    lc_img = ARCH / "1776836651782.png"
    add_image(doc, lc_img, width=COL_W)
    p_caption(doc, "표 4. Epoch별 성능 변화")
    p_body(doc,
        "학습률 감소 시점(epoch 10, 20)에서 mAP가 단계적으로 상승하며 StepLR 스케줄러의 효과를 입증한다. "
        "Epoch 27에서 최적 mAP@0.5 0.737을 달성하였다.", indent=True)

    p_heading(doc, "4.6 정성적 분석", size=9)
    p_body(doc,
        "정성적 분석은 CMAFM이 '왜 효과적인가'와 '어떤 조건에서 특히 효과적인가'를 특징맵 시각화 및 "
        "조건별 수치를 통해 함께 검증한다. 그림 2~4는 모델 처리 흐름 순서에 따라 각 단계 출력을 제시하며, "
        "표 5~6은 야간/주간 조건별 성능을 정량화한다.", indent=True)

    # 범주 표
    p_body(doc, "표. 그림 범주 설명", bold=True, size=8)
    t_leg = doc.add_table(rows=8, cols=3)
    t_leg.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t_leg, COL_W)
    leg_data = [["범주","모델 단계","설명"],
                ["(a)","입력","RGB 영상"],
                ["(b)","입력","Thermal(LWIR) 영상"],
                ["(c)","백본 출력","RGB Feature (C4)"],
                ["(d)","백본 출력","Thermal Feature (C4)"],
                ["(e)","CMAFM 출력","융합 Feature (C4)"],
                ["(f)","검출 헤드 ― 최초","RGB 단독 검출"],
                ["(g)","검출 헤드 ― 최종","RGB+Thermal 융합 검출"]]
    for i, row_data in enumerate(leg_data):
        row = t_leg.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    style_table(t_leg, font_size=7)
    for cell in t_leg.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    p_heading(doc, "4.6.1 야간 장면", size=9)
    add_image(doc, IMGS / "fig_night_comparison.png", width=COL_W)
    p_caption(doc, "그림 2. 야간 장면에서의 모델 단계별 처리 결과.\n"
              "(a) RGB 입력, (b) Thermal 입력, (c) RGB Backbone Feature (C4),\n"
              "(d) Thermal Backbone Feature (C4), (e) CMAFM 융합 Feature (C4),\n"
              "(f) RGB 단독 검출 [최초], (g) RGB+Thermal 융합 검출 [최종].")
    p_body(doc,
        "야간 장면에서 RGB Feature (c)는 조도 부족으로 활성화 강도가 낮고 객체 경계가 불분명하다. "
        "반면 Thermal Feature (d)는 열 방출 영역에 뚜렷한 활성화를 보인다. CMAFM을 거친 Fused Feature (e)는 "
        "채널 교차 어텐션이 Thermal 특징에 높은 가중치를 부여하고, 공간 교차 게이팅이 열원 위치 기준으로 "
        "RGB 특징을 선택적으로 강화하여 객체 영역에 집중된 활성화 패턴을 형성한다. "
        "(f)에서 미검출된 보행자가 (g)에서 정확히 탐지되어 융합 효과가 직접 확인된다. "
        "이는 표 5의 야간 mAP@0.5 0.853(+12.5%p), Recall 0.909(+3.9%p)로 정량적으로 뒷받침된다.", indent=True)

    p_heading(doc, "4.6.2 주간 장면", size=9)
    add_image(doc, IMGS / "fig_day_comparison.png", width=COL_W)
    p_caption(doc, "그림 3. 주간 장면에서의 모델 단계별 처리 결과.\n"
              "(a)~(g) 범주는 그림 2와 동일.")
    p_body(doc,
        "주간 장면에서 RGB Feature (c)는 텍스처·차량 외형·건물 윤곽 등 풍부한 시각 정보를 제공한다. "
        "Thermal Feature (d)는 차량 엔진부와 보행자 체열 위치에 집중된 활성화로 겹침 객체 및 소형 보행자를 "
        "보완한다. Fused Feature (e)는 게이트 가중치 α가 주간에서 RGB에 더 높은 비중을 두면서도 Thermal을 "
        "선택적으로 통합하여 객체 경계가 더 선명하고 배경 반응이 억제된다. "
        "(f) 대비 (g)에서 소형 객체 및 부분 가림(occlusion) 상황의 누락 박스가 감소한다.", indent=True)

    p_heading(doc, "4.6.3 야간/주간 조건별 정량 분석", size=9)
    p_body(doc,
        "M3FD 검증 세트(840장)를 RGB 평균 밝기 기준으로 야간(< 60, 61장)과 주간(≥ 60, 779장)으로 분류하였다.",
        indent=True)

    # 표 5
    p_body(doc, "표 5. 야간/주간 조건별 검출 성능", bold=True, size=8)
    t5 = doc.add_table(rows=4, cols=4)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t5, COL_W)
    data5 = [["조건","mAP@0.5","Recall","Miss Rate"],
             ["전체 (840)","0.737","0.874","0.126"],
             ["야간 (61)","0.853★","0.909★","0.091★"],
             ["주간 (779)","0.728","0.870","0.130"]]
    for i, row_data in enumerate(data5):
        row = t5.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    style_table(t5, font_size=8)
    for cell in t5.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for cell in t5.rows[2].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    p_body(doc,
        "야간에서 mAP@0.5 0.853으로 주간(0.728) 대비 +12.5%p, Recall +3.9%p, Miss Rate −4.0%p를 달성하였다. "
        "4.6.1의 시각화에서 확인된 바와 같이 LWIR의 열 시그니처가 CMAFM을 통해 RGB의 정보 공백을 보완하기 "
        "때문이다. 클래스별로는 대형 객체에서 야간 향상이 두드러진다(표 6).", indent=True)

    # 표 6
    p_body(doc, "표 6. 클래스별 AP@0.5 야간/주간 비교", bold=True, size=8)
    t6 = doc.add_table(rows=7, cols=4)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t6, COL_W)
    data6 = [["클래스","야간","주간","야간−주간"],
             ["Bus","0.960★","0.743","+21.7%p"],
             ["People","0.862★","0.749","+11.3%p"],
             ["Lamp","0.757★","0.670","+8.7%p"],
             ["Car","0.952★","0.867","+8.5%p"],
             ["Motorcycle","0.736★","0.695","+4.1%p"],
             ["Truck","— (없음)","0.647","—"]]
    for i, row_data in enumerate(data6):
        row = t6.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    style_table(t6, font_size=8)
    for cell in t6.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    add_image(doc, IMGS / "fig_night_vs_day.png", width=COL_W)
    p_caption(doc,
        "그림 4. 야간 vs 주간 장면에서의 RGB·Thermal·융합 특징·검출 결과 및 Thermal 기여도 비교.\n"
        "야간(상단)에서 Thermal 기여도(e열)가 주간(하단) 대비 현저히 높게 나타난다.")

    add_image(doc, IMGS / "fig_classwise_bar.png", width=COL_W)
    p_caption(doc, "그림 5. 클래스별 AP@0.5 — 전체/야간/주간 비교.")

    # ══════════════════════════════════════════════════════
    # V. 결론 및 향후 연구
    # ══════════════════════════════════════════════════════
    p_heading(doc, "V. 결론 및 향후 연구", size=10)
    p_body(doc,
        "본 논문에서는 RGB-Thermal 다중 스펙트럼 영상 융합을 위한 교차 모달 어텐션 기반 객체 검출 모델을 "
        "제안하였다. CMAFM은 채널 교차 어텐션과 공간 교차 게이팅의 이중 구조로 두 모달리티 간 전역 및 지역 "
        "수준 정보를 효율적으로 교환한다. M3FD 실험에서 mAP@0.5 73.7%, Recall 87.4%를 달성하여 RGB 단일 "
        "모달리티 대비 mAP@0.5 10.5%p, Recall 4.5%p 향상되었다.", indent=True)
    p_body(doc,
        "야간 조건에서 mAP@0.5 85.3%, Recall 90.9%로 주간 대비 각각 12.5%p, 3.9%p 높은 성능을 보였으며, "
        "Miss Rate도 4.0%p 낮아졌다. 이는 LWIR의 열 시그니처가 조도 불량 환경에서 RGB의 정보 부재를 "
        "효과적으로 보완함을 실증한다.", indent=True)
    p_body(doc,
        "본 연구 결과는 현재 우리 군이 운용 중인 감시·정찰 체계에 직접적인 기술적 기여를 할 수 있다. "
        "군에서 활용 중인 가시광 단독 객체탐지 체계는 야간·역광·연막 등 조도 불량 전술 환경에서 탐지 성능이 "
        "저하되는 한계를 갖는다. 제안 모델은 이를 세 가지 측면에서 보완한다.", indent=True)
    p_body(doc,
        "첫째, 야간 표적 탐지 능력 향상이다. 야간 Recall 90.9%(미검출률 9.1%)는 야간 침투·접근 경보 체계에 "
        "적용 시 가시광 단독 대비 현저히 낮은 미탐지율을 의미하며, 보행자 AP +11.3%p, Bus +21.7%p의 향상은 "
        "인원 및 차량 표적 식별 신뢰도를 높인다.", indent=True)
    p_body(doc,
        "둘째, 다중 환경 조건에 대한 강인성이다. CMAFM은 조도 조건에 따라 RGB와 LWIR의 기여 가중치를 자동 "
        "조정하므로 주간·야간·박명(薄明) 등 다양한 작전 시간대에 단일 모델로 일관된 탐지 성능을 유지한다.",
        indent=True)
    p_body(doc,
        "셋째, 감시·정찰 무인체계(UAV, UGV)로의 확장 가능성이다. RGB와 LWIR 카메라를 동시 탑재한 무인 "
        "감시 플랫폼에 제안 모델을 탑재하면 표적 자동 탐지·추적 신뢰성을 높일 수 있으며, 향후 경량화 연구를 "
        "통해 전방 감시 초소(GOP) 자동화 및 무인 경계 체계에 직접 적용할 수 있을 것으로 기대된다.", indent=True)

    p_heading(doc, "향후 연구 방향", size=9)
    items = [
        "검출기 고도화: Faster R-CNN을 YOLOv8 또는 RT-DETR로 교체하여 성능·속도 동시 개선",
        "조건별 학습 전략: 야간·흐림 조건 가중 샘플링 또는 condition-aware attention 도입",
        "다중 데이터셋 검증: KAIST, LLVIP 등 추가 데이터셋에서 일반화 성능 검증",
        "경량화: Knowledge Distillation 또는 백본 경량화를 통한 실시간 응용 가능성 탐색",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(9)
        run.font.name = "맑은 고딕"
        run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))

    # ══════════════════════════════════════════════════════
    # 참고문헌
    # ══════════════════════════════════════════════════════
    p_heading(doc, "참고문헌", size=10)
    refs = [
        "[1] Z. Zou et al., \"Object Detection in 20 Years: A Survey,\" Proc. IEEE, vol. 111, no. 3, pp. 257-332, 2023.",
        "[2] S. Hwang et al., \"Multispectral Pedestrian Detection: Benchmark Dataset and Baseline,\" CVPR, 2015.",
        "[3] C. Li et al., \"Multispectral Pedestrian Detection via Simultaneous Detection and Segmentation,\" BMVC, 2018.",
        "[4] J. Liu et al., \"TarDAL: Target Detection and Domain Adaptation in Multispectral Imaging,\" CVPR, 2022.",
        "[5] H. Zhang et al., \"Multispectral Fusion for Object Detection with Cyclic Fuse-and-Refine Blocks,\" ICIP, 2020.",
        "[6] F. Cao et al., \"Cross-Modal Feature Fusion for RGB-Thermal Object Detection,\" IEEE T-ITS, 2023.",
        "[7] K. Kim, \"Survey on Multispectral Pedestrian Detection,\" Journal of IEIE, vol. 60, no. 1, 2023.",
        "[8] J. Hu et al., \"Squeeze-and-Excitation Networks,\" CVPR, 2018.",
        "[9] S. Ren et al., \"Faster R-CNN: Towards Real-Time Object Detection with RPN,\" NeurIPS, 2015.",
        "[10] T.-Y. Lin et al., \"Feature Pyramid Networks for Object Detection,\" CVPR, 2017.",
        "[15] J. Liu et al., \"Target-aware Dual Adversarial Learning and M3FD Benchmark,\" CVPR, 2022.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(ref)
        run.font.size = Pt(8)
        run.font.name = "맑은 고딕"
        run._r.get_or_add_rPr().append(_eastAsia("맑은 고딕"))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"저장 완료: {OUT}")


if __name__ == "__main__":
    build_doc()
